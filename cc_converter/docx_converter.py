from pathlib import Path
from typing import Dict, Optional, List, Tuple, Union
import zipfile

from .models import Assessment, QuestionType, TextRun, TextStyle, ImageInfo, TextContent
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocxConverter:
    """Converter for Assessment objects to docx files."""

    def __init__(self, font_mapping: Optional[Dict[str, str]] = None, template_path: Optional[Path] = None):
        """Initialize the converter.

        Args:
            font_mapping: An optional mapping from font names in the source to
                font names in docx.
            template_path: Optional path to a template docx file. If not provided,
                will use the default template in the package.
        """
        self.font_mapping = font_mapping or {}
        self.template_path = template_path or Path(__file__).parent / 'template.docx'

    def convert_assessment(
        self, assessment: Assessment, output_path: Path,
        resource_zip: Optional[zipfile.ZipFile] = None,
        is_answer_key: bool = False
    ):
        """Convert an Assessment object to a docx file.

        Args:
            assessment: The Assessment object to convert.
            output_path: The path where the docx file will be saved.
            resource_zip: An optional zipfile containing resources such as images.
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "The 'python-docx' package is required to create docx files"
            )

        # Create a new document from template
        doc = Document(self.template_path)

        # Add the title
        doc.add_heading(assessment.title.strip(), level=1)

        # Process each section
        for section_idx, section in enumerate(assessment.sections, 1):
            # Process each item
            for item_idx, item in enumerate(section.items, 1):
                # Use appropriate style based on question type
                style_name = 'Question'
                question_para = doc.add_paragraph(style=style_name)
                
                # Add question text with prefix for multiple choice
                if item.question_type == QuestionType.MULTIPLE_CHOICE:
                    
                    if is_answer_key:
                        # Find index where ident is correct response
                        correct_option_idx = next((idx for idx, opt in enumerate(item.response_options) if opt.ident == item.correct_response), None)
                        if correct_option_idx is not None:
                            question_para.add_run(f"__{chr(65 + correct_option_idx)}__ ")
                    else:
                        question_para.add_run("_____ ")


                self._add_content(question_para, item.text, resource_zip)

                if item.question_type == QuestionType.MULTIPLE_CHOICE:
                    for opt_idx, option in enumerate(item.response_options):
                        option_para = doc.add_paragraph(style='MultiAns')
                        self._add_content(option_para, option.text, resource_zip)

        # Save the document
        doc.save(output_path)

    def _add_content(self, paragraph, content: List[TextContent], resource_zip: Optional[zipfile.ZipFile] = None):
        """Add text runs and images to a paragraph with proper styling."""
        for idx, item in enumerate(content):
            if isinstance(item, TextRun):
                # It's a text run
                text = item.text
                if idx == len(content) - 1:
                    text = text.rstrip('\n')
                # if len(text) > 0:
                #     print(f"Ends with {ord(text[-2:][0])}, {ord(text[-1:][0])} '{text[-10:]}'")
                style = item.style
                
                # Create a docx run
                docx_run = paragraph.add_run(text)
                
                # Apply styling
                self._apply_style_to_run(docx_run, style)
            elif isinstance(item, ImageInfo):
                # It's an inline image
                self._add_inline_image(paragraph, item, resource_zip)

    def _add_inline_image(self, paragraph, img: ImageInfo, resource_zip: Optional[zipfile.ZipFile] = None):
        """Add an inline image to the paragraph with specified dimensions."""
        from docx.shared import Inches, Mm
        
        # Check if it's a URL
        img_path = img.src
        width_param = None
        height_param = None
        
        # Convert dimensions to docx units (mm)
        if img.width:
            # Convert pixels to mm (assuming 96 DPI)
            width_param = Mm(img.width * 0.264583)
        
        if img.height:
            # Convert pixels to mm (assuming 96 DPI)
            height_param = Mm(img.height * 0.264583)
        
        if img_path.startswith('http://') or img_path.startswith('https://'):
            try:
                import requests
                from io import BytesIO
                
                # Download the image
                response = requests.get(img_path, timeout=10)
                response.raise_for_status()
                
                # Add the image as a run in the paragraph
                run = paragraph.add_run()
                if width_param and height_param:
                    run.add_picture(BytesIO(response.content), width=width_param, height=height_param)
                elif width_param:
                    run.add_picture(BytesIO(response.content), width=width_param)
                elif height_param:
                    run.add_picture(BytesIO(response.content), height=height_param)
                else:
                    run.add_picture(BytesIO(response.content))
            except Exception as e:
                # Log error but continue
                print(f"Error adding image from URL: {img_path} - {str(e)}")
        elif resource_zip:
            # It's a local path in the zip file
            try:
                with resource_zip.open(img_path) as img_file:
                    run = paragraph.add_run()
                    if width_param and height_param:
                        run.add_picture(img_file, width=width_param, height=height_param)
                    elif width_param:
                        run.add_picture(img_file, width=width_param)
                    elif height_param:
                        run.add_picture(img_file, height=height_param)
                    else:
                        run.add_picture(img_file)
            except (KeyError, zipfile.BadZipFile) as e:
                # Log error but continue
                print(f"Error adding image from zip: {img_path} - {str(e)}")
        else:
            # Log error but continue
            print(f"Error adding image: {img_path} - No resource zip provided")

    def _apply_style_to_run(self, docx_run, style: TextStyle):
        """Apply text styling to a docx run."""
        from docx.shared import Pt, RGBColor

        # Set font family if available
        if style.font_family:
            mapped_font = self.font_mapping.get(style.font_family, style.font_family)
            docx_run.font.name = mapped_font

        # Set font size if available
        if style.font_size:
            docx_run.font.size = Pt(style.font_size)

        # Set bold
        docx_run.font.bold = style.bold

        # Set color if available
        if style.color:
            # Parse color string (like "#000000" or "rgb(0,0,0)")
            rgb = self._parse_color(style.color)
            if rgb:
                docx_run.font.color.rgb = RGBColor(*rgb)

        # Set superscript/subscript
        docx_run.font.superscript = style.superscript
        docx_run.font.subscript = style.subscript

    def _parse_color(self, color_str: str) -> Optional[Tuple[int, int, int]]:
        """Parse a color string into RGB values."""
        import re

        # Handle hex color values
        if color_str.startswith("#"):
            color_hex = color_str.lstrip("#")
            try:
                if len(color_hex) == 3:  # Short form like #RGB
                    r = int(color_hex[0] + color_hex[0], 16)
                    g = int(color_hex[1] + color_hex[1], 16)
                    b = int(color_hex[2] + color_hex[2], 16)
                    return (r, g, b)
                elif len(color_hex) == 6:  # Long form like #RRGGBB
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    return (r, g, b)
            except ValueError:
                pass

        # Handle rgb() format
        rgb_match = re.match(
            r"rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)",
            color_str
        )
        if rgb_match:
            try:
                r = int(rgb_match.group(1))
                g = int(rgb_match.group(2))
                b = int(rgb_match.group(3))
                return (r, g, b)
            except ValueError:
                pass

        # Handle common color names
        color_map = {
            "black": (0, 0, 0),
            "white": (255, 255, 255),
            "red": (255, 0, 0),
            "green": (0, 128, 0),
            "blue": (0, 0, 255),
            "yellow": (255, 255, 0),
            "purple": (128, 0, 128),
            "cyan": (0, 255, 255),
            "magenta": (255, 0, 255),
            "gray": (128, 128, 128),
            "grey": (128, 128, 128),
        }

        color_str = color_str.lower()
        if color_str in color_map:
            return color_map[color_str]

        return None


def convert_assessment_to_docx(
    assessment: Assessment,
    output_path: Union[str, Path],
    resource_zip: Optional[Union[str, zipfile.ZipFile]] = None,
    font_mapping: Optional[Dict[str, str]] = None,
    is_answer_key: bool = False,
    template_path: Optional[Union[str, Path]] = None
):
    """Convert an Assessment object to a docx file.

    Args:
        assessment: The Assessment object to convert.
        output_path: The path where the docx file will be saved.
        resource_zip: An optional zipfile or path to a zipfile containing resources.
        font_mapping: An optional mapping from font names in the source to font names in docx.
        is_answer_key: Whether to generate an answer key version.
        template_path: Optional path to a template docx file. If not provided,
            will use the default template in the package.
    """
    output_path = Path(output_path)
    if template_path:
        template_path = Path(template_path)

    # Create parent directory if it doesn't exist
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Handle resource_zip parameter
    zip_to_close = None
    if isinstance(resource_zip, str):
        zip_to_close = zipfile.ZipFile(resource_zip, 'r')
        resource_zip = zip_to_close

    try:
        # Create converter and convert
        converter = DocxConverter(font_mapping, template_path)
        converter.convert_assessment(assessment, output_path, resource_zip, is_answer_key)
    finally:
        # Close the zipfile if we opened it
        if zip_to_close:
            zip_to_close.close()


def convert_cartridge_to_docx(
    cartridge_path: Union[str, Path],
    output_dir: Union[str, Path],
    font_mapping: Optional[Dict[str, str]] = None,
    limit: Optional[int] = None,
    template_path: Optional[Union[str, Path]] = None
):
    """Extract assessments from a cartridge and convert them to docx files.

    Args:
        cartridge_path: Path to the cartridge file.
        output_dir: Directory where docx files will be saved.
        font_mapping: An optional mapping from font names in the source to font names in docx.
        limit: Maximum number of assessments to process (default: all).
        template_path: Optional path to a template docx file. If not provided,
            will use the default template in the package.
    """
    from .xml_parser import parse_cartridge

    # Normalize paths
    cartridge_path = Path(cartridge_path)
    output_dir = Path(output_dir)
    if template_path:
        template_path = Path(template_path)

    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)

    # Parse the cartridge
    assessments = parse_cartridge(cartridge_path, font_mapping, limit)

    # Convert each assessment
    with zipfile.ZipFile(cartridge_path, 'r') as zf:
        for assessment in assessments:
            # Create a valid filename from the assessment title
            filename = assessment.title
            for char in '<>:"/\\|?*':
                filename = filename.replace(char, '_')

            convert_assessment_to_docx(assessment, output_dir / f"{filename}.docx", zf, font_mapping, is_answer_key=False, template_path=template_path)
            convert_assessment_to_docx(assessment, output_dir / f"{filename}_key.docx", zf, font_mapping, is_answer_key=True, template_path=template_path)

    return len(assessments) 

