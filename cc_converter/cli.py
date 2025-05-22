import argparse
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Iterable

from .models import Document, TextRun
from .parser import parse_assessment


def parse_args(argv=None):
    parser = argparse.ArgumentParser(
        description="Convert 1EdTech Common Cartridge files to doc files"
    )
    parser.add_argument("cartridge", type=Path, help="Path to .imscc file")
    parser.add_argument(
        "output", type=Path, help="Directory to write output document", nargs="?", default=Path("output")
    )
    return parser.parse_args(argv)


def _add_runs(paragraph, runs: Iterable[TextRun]) -> None:
    for run_data in runs:
        run = paragraph.add_run(run_data.text)
        run.font.superscript = run_data.superscript
        run.font.subscript = run_data.subscript


def _create_docx(doc_path: Path, document: Document, zf: zipfile.ZipFile) -> None:
    """Create a docx file for ``document`` using ``python-docx``."""

    try:
        from docx import Document as Doc
    except Exception as exc:  # pragma: no cover - raises when dependency missing
        raise ImportError(
            "The 'python-docx' package is required to create docx files"
        ) from exc

    doc = Doc()
    doc.add_heading(document.title, level=1)
    for idx, question in enumerate(document.questions, start=1):
        doc.add_paragraph(f"Question {idx}:")
        qp = doc.add_paragraph()
        _add_runs(qp, question.text)
        for img in question.images:
            try:
                with zf.open(img) as img_file:
                    doc.add_picture(img_file)
            except KeyError:
                pass
        for ans in question.answers:
            ap = doc.add_paragraph(style="List Bullet")
            _add_runs(ap, ans.text)
            for img in ans.images:
                try:
                    with zf.open(img) as img_file:
                        doc.add_picture(img_file)
                except KeyError:
                    pass

    doc.save(doc_path)


def convert_cartridge_to_doc(cartridge: Path, output_dir: Path):
    """Convert ``cartridge`` into one docx per resource in ``output_dir``."""

    output_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(cartridge, "r") as zf:
        with zf.open("imsmanifest.xml") as manifest_file:
            tree = ET.parse(manifest_file)

        root = tree.getroot()
        ns = {"ns": root.tag.split("}")[0].strip("{")}
        resources = root.findall(".//ns:resource", ns)

        for res in resources:
            file_elem = res.find("ns:file", ns)
            if file_elem is None:
                continue
            href = file_elem.get("href")
            if not href:
                continue
            with zf.open(href) as f:
                xml_text = f.read().decode("utf-8")
            doc_model = parse_assessment(xml_text)
            name = doc_model.title or res.get("identifier")
            illegal = '<>:"/\\|?*'
            for ch in illegal:
                name = name.replace(ch, "_")
            doc_path = output_dir / f"{name}.docx"
            _create_docx(doc_path, doc_model, zf)

    print(f"Created {len(resources)} docx files in {output_dir}")


def main(argv=None):
    args = parse_args(argv)
    convert_cartridge_to_doc(args.cartridge, args.output)


if __name__ == "__main__":  # pragma: no cover
    main()
